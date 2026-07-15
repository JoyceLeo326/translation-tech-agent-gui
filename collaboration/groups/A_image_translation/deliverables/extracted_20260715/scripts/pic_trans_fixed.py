from __future__ import annotations

import os
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

import pypdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
ORIGINAL_DOCX = next(p for p in sorted(ROOT.glob("*.docx")) if "已翻译" not in p.name and "完整修正版" not in p.name)
PDF_SIZE_BYTES = 10_978_986
PDF_SOURCE = next(p for p in (Path.home() / "Downloads").glob("*.pdf") if p.stat().st_size == PDF_SIZE_BYTES)

ORIGINAL_IMAGES = ROOT / "all_pictures"
RECOVERED_IMAGES = ROOT / "recovered_missing_images"
FIXED_IMAGES = ROOT / "translated_images_fixed"
OUTPUT_DOCX = ROOT / "翻译资源编写-中国文化知识百科_完整修正版.docx"
MANIFEST_XLSX = ROOT / "translation_manifest_fixed.xlsx"

FONT_REGULAR = Path(os.environ.get("PIC_TRANS_FONT_REGULAR", r"C:\Windows\Fonts\arial.ttf"))
FONT_BOLD = Path(os.environ.get("PIC_TRANS_FONT_BOLD", r"C:\Windows\Fonts\arialbd.ttf"))


@dataclass(frozen=True)
class TextBlock:
    rect: tuple[int, int, int, int]
    text: str
    max_size: int = 42
    min_size: int = 11
    fill: tuple[int, int, int] = (252, 246, 225)
    outline: tuple[int, int, int] | None = (214, 187, 132)
    bold: bool = True


def font_path(bold: bool) -> str | None:
    candidate = FONT_BOLD if bold else FONT_REGULAR
    if candidate.exists():
        return str(candidate)
    if FONT_REGULAR.exists():
        return str(FONT_REGULAR)
    return None


def get_font(size: int, bold: bool = True) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = font_path(bold)
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def fit_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    max_size: int,
    min_size: int,
    bold: bool,
) -> tuple[ImageFont.ImageFont, list[str], int]:
    x0, y0, x1, y1 = box
    width = max(1, x1 - x0 - 18)
    height = max(1, y1 - y0 - 12)
    for size in range(max_size, min_size - 1, -1):
        font = get_font(size, bold)
        lines = wrap_text(draw, text, font, width)
        bboxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
        line_heights = [max(1, b[3] - b[1]) for b in bboxes]
        total_h = sum(line_heights) + max(0, len(lines) - 1) * int(size * 0.25)
        max_w = max((b[2] - b[0] for b in bboxes), default=0)
        if total_h <= height and max_w <= width:
            return font, lines, int(size * 0.25)
    font = get_font(min_size, bold)
    return font, wrap_text(draw, text, font, width), max(2, int(min_size * 0.2))


def draw_block(draw: ImageDraw.ImageDraw, block: TextBlock) -> None:
    x0, y0, x1, y1 = block.rect
    radius = max(6, min(24, (y1 - y0) // 5))
    draw.rounded_rectangle(block.rect, radius=radius, fill=block.fill, outline=block.outline, width=2 if block.outline else 1)
    font, lines, spacing = fit_text(draw, block.text, block.rect, block.max_size, block.min_size, block.bold)
    bboxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    line_heights = [max(1, b[3] - b[1]) for b in bboxes]
    total_h = sum(line_heights) + max(0, len(lines) - 1) * spacing
    y = y0 + ((y1 - y0) - total_h) / 2
    for line, bbox, line_h in zip(lines, bboxes, line_heights):
        w = bbox[2] - bbox[0]
        x = x0 + ((x1 - x0) - w) / 2
        draw.text((x, y - bbox[1]), line, font=font, fill=(28, 28, 28))
        y += line_h + spacing


BLOCKS: dict[str, list[TextBlock]] = {
    "image2.jpeg": [
        TextBlock((500, 70, 1145, 165), "Cultivate virtue; follow nature", 46),
        TextBlock((470, 785, 1188, 865), "Confucian self-cultivation; Daoist freedom", 34),
    ],
    "image3.jpeg": [
        TextBlock((890, 100, 1585, 190), "Zen dwells among mountains and rivers; the mind rests with quiet clouds", 30),
        TextBlock((1110, 188, 1370, 245), "Words end; meaning is endless", 24),
    ],
    "image4.jpeg": [
        TextBlock((500, 382, 1140, 466), "Brush and ink carry mountains and rivers; writing shines for a thousand years", 29),
        TextBlock((632, 470, 1012, 530), "The Chinese spirit in poetry, lyrics, drama, and prose", 22),
    ],
    "image5.png": [
        TextBlock((445, 70, 1210, 170), "The grandeur of Tang; the elegance of Song", 42),
        TextBlock((525, 758, 1120, 860), "Poetry voices aspiration; lyrics convey feeling", 35),
    ],
    "image6.png": [
        TextBlock((455, 380, 1215, 530), "Every stroke carries the spirit of mountains and rivers", 40),
        TextBlock((610, 498, 1030, 555), "Calligraphy is the living art of Chinese characters", 22),
        TextBlock((365, 588, 1290, 728), "Bronze, seal, clerical, regular, running, and cursive scripts", 38),
        TextBlock((176, 320, 260, 365), "Oracle", 16),
        TextBlock((320, 320, 390, 365), "Bone", 16),
        TextBlock((455, 320, 525, 365), "Bronze", 16),
        TextBlock((598, 320, 690, 365), "Small seal", 15),
        TextBlock((758, 320, 840, 365), "Clerical", 15),
        TextBlock((914, 320, 990, 365), "Regular", 15),
        TextBlock((1065, 320, 1140, 365), "Running", 15),
        TextBlock((1240, 320, 1310, 365), "Cursive", 15),
    ],
    "image7.png": [
        TextBlock((438, 725, 1205, 820), "Mortise and tenon interlock; bricks and tiles become poetry", 38),
        TextBlock((608, 820, 1060, 882), "Ritual order and spatial aesthetics in Chinese architecture", 22),
    ],
    "image8.tiff": [
        TextBlock((955, 72, 1590, 175), "Though made by humans, it seems opened by Heaven", 34),
        TextBlock((1065, 158, 1480, 222), "In a small garden, mountains and forests hold poetry", 24),
    ],
    "image9.png": [
        TextBlock((505, 50, 1145, 145), "One region's waters; one region's homes", 40),
        TextBlock((620, 142, 1040, 195), "Dwellings hold the wisdom of Chinese life", 22),
        TextBlock((155, 750, 265, 820), "Beijing", 22),
        TextBlock((520, 748, 730, 820), "Loess cave dwellings", 21),
        TextBlock((900, 748, 1145, 820), "Huizhou-style gable walls", 19),
        TextBlock((1350, 748, 1530, 820), "Hakka tulou", 22),
    ],
    "image10.png": [
        TextBlock((478, 720, 1175, 818), "From clay to fire, porcelain gains form and grace", 36),
        TextBlock((630, 823, 1035, 882), "Chinese porcelain is renowned worldwide", 22),
    ],
    "image11.tiff": [
        TextBlock((1060, 165, 1600, 370), "Thread by thread,\nweaving the splendid Silk Road", 40),
        TextBlock((1075, 365, 1580, 445), "Silk cultivation linked Eastern and Western civilizations", 24),
    ],
    "image12.png": [
        TextBlock((325, 0, 575, 92), "Chu-Han Contention", 28),
        TextBlock((998, 92, 1562, 222), "With one storyteller's block,\ntales unfold the rise and fall of ages", 30),
        TextBlock((1090, 188, 1485, 242), "Folk performance preserves everyday life", 21),
    ],
    "image13.png": [
        TextBlock((680, 108, 950, 182), "Local Theater Stage", 26),
        TextBlock((486, 720, 1190, 820), "One stage, one story", 43),
        TextBlock((600, 820, 1070, 882), "Three hundred opera forms take root in the countryside", 22),
    ],
    "image14.png": [
        TextBlock((1424, 96, 1580, 715), "The theater stage tells tales;\nlife's many faces endure", 28),
        TextBlock((1360, 768, 1620, 908), "Yuan drama and Ming-Qing fiction:\noral storytelling lives on", 22),
    ],
    "image15.png": [
        TextBlock((704, 150, 935, 718), "The seasons turn;\ncustoms are handed down", 31),
        TextBlock((690, 470, 810, 710), "Reunion,\ngratitude,\nand respect\nfor elders", 21),
    ],
    "image16_p05_top_landscape_painting.jpg": [
        TextBlock(
            (545, 705, 1715, 1225),
            "With brush and ink, Heaven and earth appear;\nin mountains and waters, the heart's nature is seen\nUnity of Heaven and humanity; poetry lies beyond the painting",
            38,
            fill=(255, 255, 255),
            outline=None,
            bold=False,
        ),
    ],
    "image17_p05_bottom_flower_bird_painting.jpg": [
        TextBlock(
            (500, 80, 1760, 275),
            "In one flower and one leaf,\nthe scholar's spirit resides",
            48,
            fill=(255, 255, 255),
            outline=None,
        ),
        TextBlock(
            (735, 238, 1515, 325),
            "Orchids, bamboo, plum, and chrysanthemum: symbols of character",
            25,
            fill=(255, 255, 255),
            outline=None,
            bold=False,
        ),
    ],
    "image18_p07_bottom_four_inventions.jpg": [
        TextBlock(
            (485, 895, 1770, 1205),
            "Ingenious craft; objects carry civilization\nThe four great inventions advanced world history",
            52,
            fill=(255, 255, 255),
            outline=None,
            bold=False,
        ),
    ],
}


def recover_missing_images() -> None:
    RECOVERED_IMAGES.mkdir(exist_ok=True)
    reader = pypdf.PdfReader(str(PDF_SOURCE))
    items = [
        (5, 0, "image16_p05_top_landscape_painting.jpg"),
        (5, 1, "image17_p05_bottom_flower_bird_painting.jpg"),
        (7, 1, "image18_p07_bottom_four_inventions.jpg"),
    ]
    for page_no, image_index, name in items:
        image = list(reader.pages[page_no - 1].images)[image_index]
        (RECOVERED_IMAGES / name).write_bytes(image.data)


def render_fixed_images() -> None:
    FIXED_IMAGES.mkdir(exist_ok=True)
    shutil.copy2(ORIGINAL_IMAGES / "image1.jpeg", FIXED_IMAGES / "image1.jpeg")
    for src in sorted(ORIGINAL_IMAGES.iterdir()):
        if src.name == "image1.jpeg" or src.suffix.lower() not in {".jpg", ".jpeg", ".png", ".tif", ".tiff"}:
            continue
        image = Image.open(src).convert("RGB")
        draw = ImageDraw.Draw(image)
        for block in BLOCKS.get(src.name, []):
            draw_block(draw, block)
        target = FIXED_IMAGES / src.name
        if target.suffix.lower() in {".jpg", ".jpeg"}:
            image.save(target, quality=95)
        elif target.suffix.lower() in {".tif", ".tiff"}:
            image.save(target, compression="tiff_deflate")
        else:
            image.save(target)

    for src in sorted(RECOVERED_IMAGES.glob("image*.jpg")):
        image = Image.open(src).convert("RGB")
        draw = ImageDraw.Draw(image)
        for block in BLOCKS.get(src.name, []):
            draw_block(draw, block)
        image.save(FIXED_IMAGES / src.name, quality=95)


def write_manifest() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "fixed translations"
    ws.append(["image", "x0", "y0", "x1", "y1", "translation"])
    for image_name, blocks in BLOCKS.items():
        for block in blocks:
            x0, y0, x1, y1 = block.rect
            ws.append([image_name, x0, y0, x1, y1, block.text])
    wb.save(MANIFEST_XLSX)


def insert_picture_after(doc: Document, paragraph_element, image_path: Path, width_inches: float = 6.25) -> None:
    new_p = OxmlElement("w:p")
    paragraph_element.addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def paragraph_text(paragraph_element) -> str:
    return "".join(t.text or "" for t in paragraph_element.xpath(".//*[local-name()='t']"))


def replace_existing_images(doc: Document) -> None:
    translated = {p.name.lower(): p for p in FIXED_IMAGES.iterdir() if p.is_file()}
    for blip in doc.element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid or rid not in doc.part.rels:
            continue
        rel = doc.part.rels[rid]
        old_name = Path(rel.target_part.partname).name.lower()
        if old_name in translated:
            rel.target_part._blob = translated[old_name].read_bytes()


def build_docx() -> None:
    doc = Document(str(ORIGINAL_DOCX))
    replace_existing_images(doc)

    targets = [
        ("\u4e2d\u56fd\u5c71\u6c34\u753b", FIXED_IMAGES / "image16_p05_top_landscape_painting.jpg"),
        ("\u82b1\u9e1f\u753b\uff0c\u8bb2\u7a76", FIXED_IMAGES / "image17_p05_bottom_flower_bird_painting.jpg"),
        ("\u56db\u5927\u53d1\u660e\uff0c\u662f\u4e2d\u56fd\u732e\u7ed9\u4e16\u754c\u7684\u793c\u7269", FIXED_IMAGES / "image18_p07_bottom_four_inventions.jpg"),
    ]
    all_paragraphs = list(doc.element.xpath(".//*[local-name()='p']"))
    inserted: set[str] = set()
    for needle, image_path in targets:
        for paragraph in all_paragraphs:
            if needle in paragraph_text(paragraph):
                insert_picture_after(doc, paragraph, image_path)
                inserted.add(image_path.name)
                break
    missing = [image_path.name for _, image_path in targets if image_path.name not in inserted]
    if missing:
        raise RuntimeError(f"Could not insert recovered images: {missing}")

    doc.save(str(OUTPUT_DOCX))


def inspect_docx_media(docx_path: Path) -> dict[str, list[str]]:
    with zipfile.ZipFile(docx_path) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
        svg = [name for name in media if name.lower().endswith(".svg")]
    return {"media": media, "svg": svg}


def translate_svg_text(svg_bytes: bytes, translator) -> bytes:
    """Translate real SVG text nodes without Cairo/GTK rasterization.

    This handles SVG files that contain text as XML text/tspan nodes. If an SVG
    has already converted its words to vector paths, no script can recover the
    original text deterministically; that case must be rendered and OCRed.
    """

    cjk = re.compile(r"[\u3400-\u9fff]")
    root = ET.fromstring(svg_bytes)
    changed = False
    for node in root.iter():
        tag = node.tag.rsplit("}", 1)[-1].lower()
        if tag not in {"text", "tspan"}:
            continue
        if node.text and cjk.search(node.text):
            node.text = translator(node.text)
            changed = True
    if not changed:
        raise ValueError("SVG contains no editable CJK text nodes; render+OCR is required.")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def main() -> None:
    print(f"Original DOCX: {ORIGINAL_DOCX}")
    print(f"PDF source: {PDF_SOURCE}")
    print(f"Original DOCX SVG media: {inspect_docx_media(ORIGINAL_DOCX)['svg']}")
    recover_missing_images()
    render_fixed_images()
    write_manifest()
    build_docx()
    print(f"Fixed images: {FIXED_IMAGES}")
    print(f"Manifest: {MANIFEST_XLSX}")
    print(f"Output DOCX: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
