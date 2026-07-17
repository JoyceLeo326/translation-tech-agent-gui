"""
DOCX解析节点 - 使用python-docx解析DOCX文档，提取带层级标记的文本，正文按句子拆分
支持从段落、表格、页眉、页脚中提取文本
"""
import os
import re
import zipfile
from lxml import etree
from docx import Document
from langchain_core.runnables import RunnableConfig
from langgraph.runtime import Runtime
from coze_coding_utils.runtime_ctx.context import Context
from utils.file.file import FileOps
from graphs.state import DocxParseInput, DocxParseOutput


def _extract_paragraph_text(para) -> str:
    """提取段落文本"""
    return para.text.strip()


def _split_sentences(text: str) -> list:
    """按中文句号、问号、感叹号拆分为单句"""
    sentences = re.split(r'(?<=[。！？])', text)
    result = []
    for s in sentences:
        s = s.strip()
        if s:
            result.append(s)
    return result


def _extract_table_text(table) -> list:
    """从表格中提取所有单元格文本，按行排列"""
    lines = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if not text:
                continue
            # 单元格内可能有多行（如标题+正文），按换行拆分
            cell_lines = text.split("\n")
            for cl in cell_lines:
                cl = cl.strip()
                if cl:
                    lines.append(cl)
    return lines


def _arabic_to_roman(num: int) -> str:
    """阿拉伯数字转罗马数字（1-3999）"""
    if num <= 0 or num > 3999:
        return str(num)
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
    result = ''
    for i in range(len(val)):
        while num >= val[i]:
            result += syms[i]
            num -= val[i]
    return result


def _extract_header_footer_text(docx_path: str) -> list:
    """从DOCX的页眉页脚XML文件中提取文本，加[页眉]/[页脚]标记便于区分"""
    lines = []
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            for fname in zf.namelist():
                # 匹配 header*.xml 和 footer*.xml
                if re.match(r'word/(header|footer)\d+\.xml$', fname):
                    # 判断是页眉还是页脚
                    is_header = 'header' in fname
                    prefix = '[页眉]' if is_header else '[页脚]'
                    
                    # 读取并解析XML
                    xml_content = zf.read(fname)
                    root = etree.fromstring(xml_content)
                    
                    # 遍历所有<w:p>段落
                    for p_elem in root.iter(f'{{{W_NS}}}p'):
                        # 检查段落中是否有页码字段（fldChar）
                        has_page_field = False
                        for fld_elem in p_elem.iter(f'{{{W_NS}}}fldChar'):
                            has_page_field = True
                        
                        # 提取段落内所有<w:t>文本
                        texts = []
                        for t_elem in p_elem.iter(f'{{{W_NS}}}t'):
                            if t_elem.text:
                                texts.append(t_elem.text)
                        para_text = ''.join(texts).strip()
                        
                        # 如果有页码字段，只提取非数字文本（页码由Word自动处理，不需要提取）
                        if has_page_field:
                            # 移除数字（页码），只保留文本内容
                            non_digit_text = re.sub(r'\d+', '', para_text).strip()
                            if non_digit_text:
                                lines.append(f"{prefix}{non_digit_text}")
                        elif para_text:
                            # 普通段落，直接添加
                            lines.append(f"{prefix}{para_text}")
    except Exception as e:
        # 页眉页脚提取失败不影响主流程
        pass
    
    return lines


def docx_parse_node(state: DocxParseInput, config: RunnableConfig, runtime: Runtime[Context]) -> DocxParseOutput:
    """
    title: DOCX解析
    desc: 解析DOCX文档，识别一级大标题、二级小标题、正文段落，正文按句号问号感叹号拆分为单句，支持从表格中提取文本
    integrations: python-docx
    """
    ctx = runtime.context

    # 确保文件已提供
    file_docx = state.file_docx
    if file_docx is None:
        raise ValueError("缺少必填参数：file_docx（待解析的DOCX文件）")

    # 下载DOCX文件到本地临时目录
    docx_path = FileOps.save_to_local(file_docx, "source_doc.docx")

    # 使用python-docx解析文档
    doc = Document(docx_path)

    # 按文档顺序遍历所有元素（段落和表格）
    parsed_lines = []
    table_index = 0
    paragraph_index = 0

    # 获取文档body元素（doc.element是document，第一个子元素是body）
    body_element = doc.element[0]

    # 注意：不再需要检测默认字体，使用基于内容特征的通用标题识别
    paragraph_index = 0

    for child in body_element.iterchildren():
        tag = etree.QName(child).localname

        if tag == 'p':
            # 段落
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                paragraph_index += 1
                text = _extract_paragraph_text(para)
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"
                style_name_lower = style_name.lower() if style_name else ""

                # 通用标题识别：基于内容特征（不依赖字体、字号等格式）
                # 标题特征：短段落 + 无句号/问号/感叹号 + 无逗号
                is_short = len(text) < 30
                has_terminator = any(c in text for c in '。！？')
                has_comma = '，' in text
                
                if style_name_lower.startswith("heading 1"):
                    # 一级标题：不按句子拆分，保持完整
                    parsed_lines.append(f"[TITLE]{text}")
                elif style_name_lower.startswith("heading 2"):
                    # 二级标题：不按句子拆分，保持完整
                    parsed_lines.append(f"[TITLE]{text}")
                elif style_name_lower.startswith("heading"):
                    # 其他标题：不按句子拆分，保持完整
                    parsed_lines.append(f"[TITLE]{text}")
                elif is_short and not has_terminator and not has_comma:
                    # 通用标题识别：短段落且无标点，视为标题
                    parsed_lines.append(f"[TITLE]{text}")
                else:
                    # 正文段落：按句子拆分
                    parsed_lines.extend(_split_sentences(text))

        elif tag == 'tbl':
            # 表格
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_index += 1
                table_lines = _extract_table_text(table)
                for line in table_lines:
                    # 表格中的文本也按句子拆分（正文处理）
                    parsed_lines.extend(_split_sentences(line))

    # 后处理：合并被错误拆分的行（如"褒"和"衣广袖"应为"褒衣广袖"）
    merged_lines = []
    for line in parsed_lines:
        if merged_lines:
            prev = merged_lines[-1]
            # 获取前一行最后一个中文字符
            prev_chars = [c for c in prev if '\u4e00' <= c <= '\u9fff']
            # 获取当前行第一个中文字符
            curr_chars = [c for c in line if '\u4e00' <= c <= '\u9fff']
            # 当前行第一个字符是中文字符
            first_char = line.strip()[0] if line.strip() else ''
            is_chinese_start = '\u4e00' <= first_char <= '\u9fff'
            # 前一行不以数字开头（排除章节标题如"9.1"）
            prev_starts_with_digit = prev.strip() and prev.strip()[0].isdigit()
            # 前一行最后有中文字符（说明句子未结束）
            prev_ends_with_chinese = bool(prev_chars)
            # 检查前一行是否以句号/问号/感叹号结尾（句子已完整结束）
            prev_ends_with_terminator = prev.strip() and prev.strip()[-1] in '。！？'
            # 检查当前行是否是章节标题（如"第十章"、"10.1"、"前言"、"后记"）
            curr_is_chapter = bool(re.match(r'^(第[一二三四五六七八九十百千]+章|前言|后记|\d+\.\d+)', line.strip()))
            # 检查前一行是否是章节标题（如"后记"、"前言"）
            prev_is_chapter = bool(re.match(r'^(第[一二三四五六七八九十百千]+章|前言|后记|\d+\.\d+)', prev.strip()))
            # 检查前一行是否是标题（加粗或标题样式）
            prev_is_title = prev.startswith('[TITLE]') or prev_is_chapter
            # 合并条件：当前行以中文开头，前一行不以数字开头且以中文结尾
            # 且前一行不以句号结尾，且当前行不是章节标题，且前一行不是标题
            if is_chinese_start and not prev_starts_with_digit and prev_ends_with_chinese and not prev_ends_with_terminator and not curr_is_chapter and not prev_is_title:
                merged_lines[-1] = prev + line
                continue
        merged_lines.append(line)
    parsed_lines = merged_lines

    # 提取页眉页脚文本（加[页眉]/[页脚]标记）
    header_footer_lines = _extract_header_footer_text(docx_path)
    parsed_lines.extend(header_footer_lines)

    # 清理临时文件
    try:
        os.remove(docx_path)
    except Exception:
        pass

    parsed_text = "\n".join(parsed_lines)

    return DocxParseOutput(parsed_text=parsed_text)