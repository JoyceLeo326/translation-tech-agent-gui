"""
DOCX解析节点 - 使用python-docx解析DOCX文档，提取带层级标记的文本，正文按句子拆分
支持从段落和表格中提取文本
"""
import os
import re
from lxml import etree
from docx import Document
from langchain_core.runnables import RunnableConfig
from langgraph.runtime import Runtime
from coze_coding_utils.runtime_ctx.context import Context
from utils.file.file import FileOps
from graphs.state import DocxParseInput, DocxParseOutput


def _extract_paragraph_text(para) -> str:
    """提取段落文本"""
    return para.text.strip()


def _split_sentences(text: str) -> list:
    """按中文句号、问号、感叹号拆分为单句"""
    sentences = re.split(r'(?<=[。！？])', text)
    result = []
    for s in sentences:
        s = s.strip()
        if s:
            result.append(s)
    return result


def _extract_table_text(table) -> list:
    """从表格中提取所有单元格文本，按行排列"""
    lines = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if not text:
                continue
            # 单元格内可能有多行（如标题+正文），按换行拆分
            cell_lines = text.split("\n")
            for cl in cell_lines:
                cl = cl.strip()
                if cl:
                    lines.append(cl)
    return lines


def docx_parse_node(state: DocxParseInput, config: RunnableConfig, runtime: Runtime[Context]) -> DocxParseOutput:
    """
    title: DOCX解析
    desc: 解析DOCX文档，识别一级大标题、二级小标题、正文段落，正文按句号问号感叹号拆分为单句，支持从表格中提取文本
    integrations: python-docx
    """
    ctx = runtime.context

    # 确保文件已提供
    file_docx = state.file_docx
    if file_docx is None:
        raise ValueError("缺少必填参数：file_docx（待解析的DOCX文件）")

    # 下载DOCX文件到本地临时目录
    docx_path = FileOps.save_to_local(file_docx, "source_doc.docx")

    # 使用python-docx解析文档
    doc = Document(docx_path)

    # 按文档顺序遍历所有元素（段落和表格）
    parsed_lines = []
    table_index = 0
    paragraph_index = 0

    # 获取文档body元素（doc.element是document，第一个子元素是body）
    body_element = doc.element[0]

    for child in body_element.iterchildren():
        tag = etree.QName(child).localname

        if tag == 'p':
            # 段落
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                paragraph_index += 1
                text = _extract_paragraph_text(para)
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"
                style_name_lower = style_name.lower() if style_name else ""

                if style_name_lower.startswith("heading 1"):
                    parsed_lines.append(f"#一级标题#{text}")
                elif style_name_lower.startswith("heading 2"):
                    parsed_lines.append(f"#二级标题#{text}")
                elif style_name_lower.startswith("heading"):
                    parsed_lines.append(f"#一级标题#{text}")
                else:
                    # 正文段落：按句子拆分
                    parsed_lines.extend(_split_sentences(text))

        elif tag == 'tbl':
            # 表格
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_index += 1
                table_lines = _extract_table_text(table)
                for line in table_lines:
                    # 表格中的文本也按句子拆分（正文处理）
                    parsed_lines.extend(_split_sentences(line))

    # 后处理：合并被错误拆分的行（如"褒"和"衣广袖"应为"褒衣广袖"）
    merged_lines = []
    for line in parsed_lines:
        if merged_lines:
            prev = merged_lines[-1]
            # 获取前一行最后一个中文字符
            prev_chars = [c for c in prev if '\u4e00' <= c <= '\u9fff']
            # 获取当前行第一个中文字符
            curr_chars = [c for c in line if '\u4e00' <= c <= '\u9fff']
            # 当前行第一个字符是中文字符
            first_char = line.strip()[0] if line.strip() else ''
            is_chinese_start = '\u4e00' <= first_char <= '\u9fff'
            # 前一行不以数字开头（排除章节标题如"9.1"）
            prev_starts_with_digit = prev.strip() and prev.strip()[0].isdigit()
            # 前一行最后有中文字符（说明句子未结束）
            prev_ends_with_chinese = bool(prev_chars)
            # 检查前一行是否以句号/问号/感叹号结尾（句子已完整结束）
            prev_ends_with_terminator = prev.strip() and prev.strip()[-1] in '。！？'
            # 检查当前行是否是章节标题（如"第十章"、"10.1"、"前言"、"后记"）
            curr_is_chapter = bool(re.match(r'^(第[一二三四五六七八九十百千]+章|前言|后记|\d+\.\d+)', line.strip()))
            # 合并条件：当前行以中文开头，前一行不以数字开头且以中文结尾
            # 且前一行不以句号结尾，且当前行不是章节标题
            if is_chinese_start and not prev_starts_with_digit and prev_ends_with_chinese and not prev_ends_with_terminator and not curr_is_chapter:
                merged_lines[-1] = prev + line
                continue
        merged_lines.append(line)
    parsed_lines = merged_lines

    # 清理临时文件
    try:
        os.remove(docx_path)
    except Exception:
        pass

    parsed_text = "\n".join(parsed_lines)

    return DocxParseOutput(parsed_text=parsed_text)