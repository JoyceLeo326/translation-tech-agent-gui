from __future__ import annotations

import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from agent_gui_starter.production import (
    extract_docx_to_review,
    load_review_rows,
    refill_docx_from_review,
    synthesize_audio_from_review,
    translate_review_workbook,
)


class ProductionWorkflowTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)
        self.source = self.root / "文化故事.docx"
        document = Document()
        document.add_heading("端午节", level=1)
        document.add_paragraph("孩子们把香囊挂在胸前。大家一起看龙舟。")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "粽子香甜可口。"
        document.save(self.source)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_extract_translate_and_refill_docx(self) -> None:
        review = self.root / "review.xlsx"
        extract = extract_docx_to_review(self.source, review)
        self.assertEqual(extract.artifacts, (review,))
        rows = load_review_rows(review)
        self.assertGreaterEqual(len(rows), 4)

        translated = self.root / "translated.xlsx"
        translate_review_workbook(
            review,
            lambda lines: [f"English {index}" for index, _ in enumerate(lines, start=1)],
            translated,
            batch_size=2,
        )
        translated_rows = load_review_rows(translated)
        self.assertTrue(all(row.machine_translation.startswith("English") for row in translated_rows))

        output = self.root / "refilled.docx"
        result = refill_docx_from_review(self.source, translated, output)
        self.assertTrue(output.exists())
        report = json.loads(result.artifacts[1].read_text(encoding="utf-8"))
        self.assertEqual(report["failed_xml_files"], [])
        self.assertGreaterEqual(report["matched_items"], 3)
        with zipfile.ZipFile(output) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("English", xml)

    def test_review_workbook_has_editable_production_columns(self) -> None:
        review = self.root / "review.xlsx"
        extract_docx_to_review(self.source, review)
        workbook = load_workbook(review)
        sheet = workbook["翻译审校"]
        headers = [cell.value for cell in sheet[1]]
        self.assertEqual(
            headers,
            ["序号", "位置", "中文原文", "机器英文译文", "人工审核", "审校状态", "备注"],
        )
        self.assertEqual(sheet.freeze_panes, "A2")
        self.assertIn("使用说明", workbook.sheetnames)

    @unittest.skipUnless(__import__("sys").platform == "win32", "Windows SAPI is required")
    def test_review_workbook_generates_playable_audio_and_qr(self) -> None:
        review = self.root / "review.xlsx"
        extract_docx_to_review(self.source, review)
        workbook = load_workbook(review)
        sheet = workbook["翻译审校"]
        for row in range(2, sheet.max_row + 1):
            sheet.cell(row, 4, "A short cultural story for children.")
        workbook.save(review)
        output = self.root / "voice.wav"
        result = synthesize_audio_from_review(review, output)
        self.assertGreater(output.stat().st_size, 1024)
        self.assertTrue(result.artifacts[1].exists())
        self.assertTrue(result.artifacts[2].exists())


if __name__ == "__main__":
    unittest.main()
